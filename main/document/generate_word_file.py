from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


def generate_orcamento_doc(orcamento_id, data) -> Document:
    orcamento = data['orcamento']
    document = Document()

    heading = document.add_heading(
        f"{orcamento['evento']['nome']} - {formatar_data(orcamento['evento']['data_inicio'])}",
        level=1)
    heading.runs[0].bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph(
        f"Data: {formatar_data(orcamento['evento']['data_inicio'])} a {formatar_data(orcamento['evento']['data_fim'])}")
    document.add_paragraph(f"Horário: 08:00 às 18:00")
    document.add_paragraph(f"Local: {orcamento['evento']['local']} - {orcamento['evento'].get('descricao', '')}")

    for data_evento, intervalos in data['data'].items():
        data_formatada = formatar_data(data_evento)
        document.add_heading(f"Cardápio - {data_formatada}", level=3).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for nome_intervalo, intervalo in intervalos.items():
            if intervalo['comidas']:
                document.add_heading(nome_intervalo, level=4)
            for comida in intervalo['comidas']:
                quantidade = next(
                    (item['quantidade'] for item in orcamento['comidas']
                     if item['comida_id'] == comida['comida_id']),
                    0
                )
                nome = comida['nome'].split('-')[1]
                document.add_paragraph(
                    f"- {nome} com {quantidade} unidades"
                )
    document.add_heading(f"Equipe", level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for logistica in data['orcamento']['logisticas']:
        document.add_paragraph(f"{logistica['quantidade']} {logistica['logistica']} para {logistica['dias']} dias")
    if orcamento.get('observacoes'):
        document.add_heading("Observações", level=2)
        document.add_paragraph(orcamento['observacoes'])
    document = generate_table_values(orcamento, document)
    return document


def generate_table_values(orcamento, document: Document) -> Document:
    heading = document.add_heading("INVESTIMENTO TOTAL", level=2)
    heading.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    def set_cell_text_and_bold(cell, text):
        cell.text = ''
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True

    # Linha 1: Itens de Cardápio
    set_cell_text_and_bold(table.cell(0, 0), "Itens de Cardápio")
    valor_comidas = float(orcamento['valor_total_comidas'])
    set_cell_text_and_bold(
        table.cell(0, 1),
        f"R$ {valor_comidas:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    )

    # Linha 2: Logística/Hotel
    set_cell_text_and_bold(table.cell(1, 0), "Itens de Logística")
    valor_logistica = float(orcamento['valor_total_logisticas'])
    set_cell_text_and_bold(
        table.cell(1, 1),
        f"R$ {valor_logistica:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    )
    # Linha 3: Decoração
    set_cell_text_and_bold(table.cell(2, 0), "Total da Decoração")
    valor_decoracao = float(orcamento.get('valor_decoracao', '0.00'))
    set_cell_text_and_bold(
        table.cell(2, 1),
        f"R$ {valor_decoracao:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    )

    # Linha 4: Impostos
    set_cell_text_and_bold(table.cell(3, 0), "Impostos")
    valor_imposto = float(orcamento['valor_imposto'])
    set_cell_text_and_bold(
        table.cell(3, 1),
        f"R$ {valor_imposto:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    )

    #Linha 5: Total
    set_cell_text_and_bold(table.cell(4, 0), "Total")
    valor_total = float(orcamento['valor_total'])
    set_cell_text_and_bold(
        table.cell(4, 1),
        f"R$ {valor_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    )
    return document


def format_line_red(document: Document, text: str) -> None:
    run = document.add_paragraph().add_run(text)
    font = run.font
    font.color.rgb = RGBColor(255, 0, 0)


def formatar_data(data_str):
    try:
        return datetime.strptime(data_str, '%d-%m-%Y').strftime('%d/%m/%Y')
    except ValueError:
        try:
            return datetime.strptime(data_str, '%Y-%m-%d').strftime('%d/%m/%Y')
        except ValueError:
            return data_str
